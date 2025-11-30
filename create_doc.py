from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create document
doc = Document()

# Title
title = doc.add_heading('AI Learning Support Prompts - Impact Dashboard Project', 0)

# Introduction
intro = doc.add_paragraph('This document lists the AI prompts used for ')
intro.add_run('learning support only').bold = True
intro.add_run(' during the development of the Impact Dashboard project. All final code, design decisions, and implementations were my own work.')
doc.add_paragraph()

# Sections with prompts
sections = [
    ('React Syntax & Patterns', [
        "What's the correct syntax for using useState with multiple state variables in React?",
        "How do I properly structure a useEffect hook that runs only once on component mount?",
        "What's the difference between inline styles and CSS modules in React?",
        "Show me the pattern for conditional rendering in JSX",
        "How do I pass props between parent and child components in React?"
    ]),
    ('JavaScript & ES6', [
        "Remind me of the syntax for array destructuring in JavaScript",
        "What's the difference between map() and forEach() for rendering lists?",
        "How do I use template literals for dynamic strings?",
        "What's the correct way to use the spread operator with objects?",
        "Explain arrow function syntax vs regular functions"
    ]),
    ('Styling & UI', [
        "What are common CSS properties for creating card layouts?",
        "How do I center elements using flexbox?",
        "What's the syntax for CSS grid with responsive columns?",
        "Suggest color palette ideas for an environmental/sustainability theme",
        "What are best practices for responsive design breakpoints?"
    ]),
    ('Data Visualization', [
        "What props does Recharts AreaChart component accept?",
        "How do I format tooltip content in Recharts?",
        "What's the structure for chart data in Recharts?",
        "How do I create gradient fills in SVG charts?"
    ]),
    ('Project Structure', [
        "What's a typical folder structure for a React app?",
        "How should I organize API calls in a React project?",
        "What files should be in .gitignore for a React project?",
        "What's the purpose of package.json vs package-lock.json?"
    ]),
    ('Git & Version Control', [
        "What's the command to create a new branch in git?",
        "How do I check which branch I'm currently on?",
        "What's the difference between git add and git commit?",
        "How do I push changes to a specific branch?"
    ]),
    ('Deployment', [
        "What are the steps to deploy a React app to Vercel?",
        "How do environment variables work in React apps?",
        "What's the build command for Create React App?",
        "How do I set up a custom domain on Vercel?"
    ]),
    ('Backend Concepts', [
        "What's the basic structure of a FastAPI application?",
        "How do I define API endpoints in FastAPI?",
        "What's the purpose of CORS in web applications?",
        "How do I structure a requirements.txt file for Python?"
    ]),
    ('Debugging & Testing', [
        "What does 'failed to fetch' error mean in React?",
        "How do I check if my API is running correctly?",
        "What's the syntax for console.log debugging in React?",
        "How do I view network requests in browser DevTools?"
    ]),
    ('Acceptance Criteria Brainstorming', [
        "What should I test when implementing a navigation system?",
        "What are important user interactions to consider for a dashboard?",
        "What accessibility features should a web dashboard have?",
        "What performance metrics matter for a React application?"
    ]),
    ('Wording & Documentation', [
        "Suggest alternative phrasings for 'carbon footprint reduction'",
        "What's a professional way to describe sustainability metrics?",
        "How should I word commit messages for feature additions?",
        "What sections should a README file include?"
    ]),
    ('Learning Resources', [
        "Where can I find React documentation for hooks?",
        "What's the official Recharts documentation URL?",
        "Are there any tutorials on responsive dashboard design?"
    ])
]

counter = 1
for section_name, prompts in sections:
    doc.add_heading(section_name, 2)
    for prompt in prompts:
        doc.add_paragraph(f'{counter}. "{prompt}"')
        counter += 1
    doc.add_paragraph()

# Declaration section
doc.add_page_break()
doc.add_heading('Declaration', 1)

decl = doc.add_paragraph('All prompts above were used ')
decl.add_run('solely for learning and reference purposes').bold = True
decl.add_run('. The actual implementation, including:')

bullets = [
    'Component architecture and design decisions',
    'State management logic',
    'UI/UX design choices',
    'Data structure design',
    'Code organization',
    'Feature implementation',
    'Styling and layout'
]

for item in bullets:
    doc.add_paragraph(item, style='List Bullet')

final = doc.add_paragraph('...were all ')
final.add_run('my own original work').bold = True
final.add_run('. I did not copy-paste AI-generated code as final submissions. Any code snippets from AI responses were used only as learning references to understand syntax or patterns, which I then implemented myself according to my project requirements.')

# Save
doc.save('AI_LEARNING_SUPPORT.docx')
print('Word document created successfully!')
